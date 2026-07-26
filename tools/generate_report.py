import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    document = Document()
    
    # Set Normal text to Arial, black
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)

    # Set Title and Headings to black
    title_style = document.styles['Title']
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.font.name = 'Arial'
    title_style.font.bold = True

    for i in range(1, 4):
        heading_style = document.styles[f'Heading {i}']
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        heading_style.font.name = 'Arial'
        heading_style.font.bold = True

    # Title
    title = document.add_heading('Advanced AI Medical Intelligence Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = document.add_paragraph('Project Technical Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    
    document.add_paragraph() # Spacing

    document.add_heading('1. Project Overview', level=1)
    document.add_paragraph(
        "The Advanced AI Medical Intelligence Platform is a comprehensive, multi-modal diagnostic system "
        "designed to assist medical professionals by leveraging deep learning for image classification. "
        "The platform currently supports three core modules: Chest X-Ray Pneumonia Detection, Brain Tumor "
        "MRI Classification, and Skin Dermoscopy Cancer Detection. It provides end-to-end functionality, "
        "from image validation and AI inference to Explainable AI (XAI) visualizations and automated clinical reporting."
    )
    
    document.add_heading('2. System Architecture', level=1)
    document.add_paragraph(
        "The system follows a modern decoupled architecture, combining a high-performance backend with a reactive frontend UI."
    )
    
    document.add_heading('2.1 Backend (FastAPI & PyTorch)', level=2)
    document.add_paragraph(
        "The backend is built with FastAPI, providing asynchronous REST API endpoints. It manages image uploads, "
        "data validation, database interactions via SQLite, and coordinates the machine learning pipeline. "
        "Inference is powered by PyTorch using fine-tuned DenseNet121 models for highly accurate classifications."
    )
    
    document.add_heading('2.2 Frontend (React & Vite)', level=2)
    document.add_paragraph(
        "The user interface is a unified Single Page Application (SPA) built with React and Vite. It features a clean, "
        "professional light theme designed for medical environments. Users can select the scan type via a dropdown menu "
        "and seamlessly drag-and-drop images for immediate analysis. The UI provides visual probability bars and interactive results."
    )
    
    document.add_heading('3. Explainable AI (XAI) & Grad-CAM', level=1)
    document.add_paragraph(
        "To ensure transparency in AI decision-making, the platform implements Gradient-weighted Class Activation Mapping (Grad-CAM). "
        "For every prediction, a heatmap is generated and superimposed on the original scan. This highlights the specific anatomical "
        "regions that the DenseNet121 model focused on to make its classification, allowing clinicians to verify the AI's reasoning."
    )
    
    document.add_heading('4. Automated Clinical Reporting (xAI Grok API)', level=1)
    document.add_paragraph(
        "A key feature of the platform is the automated generation of structured clinical reports. The system integrates the advanced "
        "xAI Grok API to process diagnostic results, instead of older standard APIs like OpenAI. The Grok model interprets the predicted class, "
        "confidence scores, and modality context to draft a professional medical summary. This includes clinical findings, "
        "Grad-CAM interpretations, and recommended next steps."
    )
    
    document.add_heading('5. Deployment & Operations', level=1)
    document.add_paragraph(
        "The platform is fully containerized using Docker and Docker Compose, allowing the FastAPI backend, React frontend, "
        "and SQLite database to be launched simultaneously with a single command. The system includes health-check polling "
        "and graceful error handling for invalid image uploads or offline services."
    )
    
    document.add_heading('6. Conclusion', level=1)
    document.add_paragraph(
        "The Advanced AI Medical Intelligence Platform demonstrates how modern machine learning, explainable AI, and cutting-edge "
        "Large Language Models (xAI Grok) can be integrated into a single, cohesive clinical tool. Its unified, light-themed user "
        "interface ensures ease of use while delivering powerful diagnostic insights."
    )
    
    # Save the document
    document.save('d:\\medical-ai-platform\\Advanced AI Medical Intelligence Platform.docx')
    print("Report generated successfully.")

if __name__ == "__main__":
    create_report()
